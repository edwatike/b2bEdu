from __future__ import annotations

from enum import Enum
from typing import List, Optional, Tuple
import sys
import os
import hashlib
import httpx
import subprocess
import logging
import json
import time
import tempfile
from pathlib import Path

from pydantic import BaseModel, Field, ValidationError


logger = logging.getLogger(__name__)

try:
    from app.config import settings
except Exception:
    class _FallbackSettings:
        llm_keys_enabled: bool = (os.getenv("LLM_KEYS_ENABLED", "").strip().lower() in {"1", "true", "yes", "on"})
        llm_keys_force: bool = (os.getenv("LLM_KEYS_FORCE", "").strip().lower() in {"1", "true", "yes", "on"})
        ollama_url: str = os.getenv("OLLAMA_URL", "http://127.0.0.1:11434")
        ollama_model: str = os.getenv("OLLAMA_MODEL", "")
        ollama_timeout_sec: int = int(os.getenv("OLLAMA_TIMEOUT_SEC", "15"))

    settings = _FallbackSettings()

# Add parser_service to path for OCR modules
_repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))
_parser_service_path = os.path.join(_repo_root, "parser_service")
if os.path.isdir(_parser_service_path) and _parser_service_path not in sys.path:
    sys.path.append(_parser_service_path)

try:
    from src.simple_ocr_wrapper import smart_extract_text
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

try:
    from src.paddle_ocr_wrapper import extract_tables_from_pdf, tables_to_text, is_digital_pdf as paddle_is_digital_pdf
    PADDLEOCR_AVAILABLE = True
except ImportError:
    PADDLEOCR_AVAILABLE = False

try:
    from docling.document_converter import DocumentConverter
    DOCLING_AVAILABLE = True
except Exception:
    DOCLING_AVAILABLE = False


class RecognitionEngine(str, Enum):
    auto = "auto"
    structured = "structured"
    ocr = "ocr"
    docling = "docling"


class RecognitionDependencyError(RuntimeError):
    pass


def normalize_item_names(items: List[str]) -> List[str]:
    import re

    out: List[str] = []
    seen: set[str] = set()

    email_re = re.compile(r"\b[^\s@]+@[^\s@]+\.[^\s@]+\b", re.IGNORECASE)
    url_re = re.compile(r"\bhttps?://\S+\b|\bwww\.\S+\b", re.IGNORECASE)
    inn_kpp_re = re.compile(r"\b(инн|кпп|огрн|р/с|к/с|бик)\b", re.IGNORECASE)
    money_re = re.compile(r"\b\d{1,3}(?:[\s\u00A0]\d{3})+(?:[\.,]\d{2})?\b")
    many_digits_re = re.compile(r"\d")
    alpha_re = re.compile(r"[A-Za-zА-Яа-я]")

    currency_re = re.compile(r"\b(руб\.?|р\.?|₽|eur|usd|\$|€)\b", re.IGNORECASE)
    percent_re = re.compile(r"\b\d{1,3}(?:[\.,]\d+)?\s*%\b")
    nds_re = re.compile(r"\bндс\b", re.IGNORECASE)

    header_like_re = re.compile(
        r"\b(коммерческ(ое|ий)\s+предложени|предложени(е|я)\s+на|счет\b|сч[её]т\s*№|итого\b|всего\b|к\s+оплате|ндс\b|условия\b|контакты\b)\b",
        re.IGNORECASE,
    )
    unit_re = re.compile(r"\b(шт|шту?к|м2|м²|кг|г|т|тонн|п\.м|пог\.\s*м|м|мм|см)\b", re.IGNORECASE)
    trailing_num_re = re.compile(r"\b\d+(?:[\.,]\d+)?\b")
    # Prices and amounts often contain currency signs/words or formatted numbers.
    price_token_re = re.compile(r"^(?:\d{1,3}(?:[\s\u00A0]\d{3})+(?:[\.,]\d+)?|\d+(?:[\.,]\d+){1,2})$")

    tech_marker_re = re.compile(
        r"\b(?:dnid|dn|sn|pn|sdr|od|id|ø|d=|l=|len=|length=|ширина|длина|толщин\w*|внутр\w*|наруж\w*)\b",
        re.IGNORECASE,
    )

    tech_kv_re = re.compile(
        r"\b(?:dnid\s*\d+|dn\s*\d+|sn\s*\d+|pn\s*\d+|sdr\s*\d+|l\s*=\s*\d+(?:[\.,]\d+)?|d\s*=\s*\d+(?:[\.,]\d+)?)\b",
        re.IGNORECASE,
    )

    def _trim_tail_columns(s: str) -> str:
        # Remove typical table tails: quantities/weights/prices after the name.
        # Keep dimension-like chunks inside the name: patterns containing '*', '/', 'x', 'х'.
        parts = s.split()
        if len(parts) < 3:
            return s

        def _is_dim(token: str) -> bool:
            t = token.lower()
            return any(ch in t for ch in ["*", "/", "x", "х"])

        # If there are explicit tech markers (DN/SN/PN/DNID/L=/D= etc.), drop everything from the first marker onward.
        for i, tok in enumerate(parts):
            t = tok.strip()
            if i >= 2 and (tech_marker_re.fullmatch(t) or tech_kv_re.search(t)):
                return " ".join(parts[:i]).strip()

        # If there's a unit token, drop everything from the first unit onward.
        # Additionally, if right before the unit there is a pure number (qty), drop it too.
        for i, tok in enumerate(parts):
            if unit_re.fullmatch(tok) and i >= 2:
                cut = i
                if i - 1 >= 2 and trailing_num_re.fullmatch(parts[i - 1] or ""):
                    cut = i - 1
                return " ".join(parts[:cut]).strip()

        def _is_mark_pair(prev: str, cur: str) -> bool:
            # Preserve patterns like "А 60" / "A 60" / "класс 60" where the number is part of marking.
            if not trailing_num_re.fullmatch(cur):
                return False
            p = (prev or "").strip()
            if len(p) == 1 and p.isalpha():
                return True
            return False

        def _is_tail_column_token(tok: str) -> bool:
            t = tok.strip()
            if not t:
                return False
            if currency_re.fullmatch(t) or currency_re.search(t):
                return True
            if percent_re.fullmatch(t) or nds_re.fullmatch(t):
                return True
            if unit_re.fullmatch(t):
                return True
            if money_re.fullmatch(t):
                return True
            if price_token_re.fullmatch(t):
                return True
            # pure numbers (qty) are tail candidates
            if trailing_num_re.fullmatch(t):
                return True
            return False

        # Drop tail numeric/price/unit columns. Stop if last token is a dimension or a marking pair.
        while len(parts) >= 2:
            last = parts[-1]
            prev = parts[-2]
            if _is_dim(last):
                break
            if _is_mark_pair(prev, last):
                break
            # percent/spec columns are almost always tails
            if percent_re.fullmatch(last) or percent_re.search(last):
                parts.pop()
                continue
            if _is_tail_column_token(last):
                parts.pop()
                continue
            break
        return " ".join(parts).strip()

    def _strip_prefix(line: str) -> str:
        s = (line or "").strip()
        s = re.sub(r"^\s*№\s*\d+\s*", "", s)
        s = re.sub(r"^\s*\d{1,4}\s+[\)\.]\s*", "", s)
        s = re.sub(r"^\s*\d{1,4}\s+", "", s)
        return s.strip()

    def _looks_like_phone(line: str) -> bool:
        s = (line or "").strip()
        if not s:
            return False
        ls = s.lower()
        if "тел" in ls or "тел." in ls or "телефон" in ls:
            return True
        # Be conservative: only treat as phone when it starts with '+' or contains typical phone punctuation.
        if not (s.startswith("+") or "(" in s or ")" in s):
            return False
        digits = re.sub(r"\D", "", s)
        # Typical phone digit counts (RU + intl)
        if 9 <= len(digits) <= 15:
            return True
        return False

    for raw in items or []:
        s = " ".join(str(raw or "").split())
        if not s:
            continue
        if header_like_re.search(s) or s.lower().startswith("предложение "):
            continue
        if email_re.search(s) or url_re.search(s) or _looks_like_phone(s) or inn_kpp_re.search(s):
            continue
        # Drop lines that look like full invoice rows with money totals
        if money_re.search(s):
            # keep only left part before first big money token
            m = money_re.search(s)
            if m:
                s = s[: m.start()].strip()
        # Also trim at first explicit currency marker, if present
        m2 = currency_re.search(s)
        if m2:
            s = s[: m2.start()].strip()
        s = _strip_prefix(s)
        s = _trim_tail_columns(s)
        # Drop extremely short/mostly-numeric strings
        digits = len(many_digits_re.findall(s))
        letters = len(alpha_re.findall(s))
        if len(s) < 4:
            continue
        # If line still looks like spec-dump, try to salvage the left "name" part instead of dropping.
        if digits >= 18 and letters >= 10 and len(s) > 40:
            # keep only the first 3-7 tokens that contain letters
            toks = [t for t in s.split() if t and alpha_re.search(t)]
            if len(toks) >= 3:
                s = " ".join(toks[:7]).strip()
                digits = len(many_digits_re.findall(s))
                letters = len(alpha_re.findall(s))
        # Drop lines that still look like spec-dumps after salvage
        if digits >= 18 and letters < 10 and len(s) > 40:
            continue
        # Only drop if it's mostly digits and has almost no letters (table debris)
        if digits > 25 and len(s) < 80 and letters < 3:
            continue
        low = s.lower()
        if low in seen:
            continue
        seen.add(low)
        out.append(s)

    return out


def extract_item_names_via_groq(*, text: str, api_key: str) -> List[str]:
    names, _usage = extract_item_names_via_groq_with_usage(text=text, api_key=api_key)
    return names


def extract_text_best_effort(*, filename: str, content: bytes, engine: RecognitionEngine = RecognitionEngine.auto) -> str:
    name = (filename or "").lower()
    ext = name.split(".")[-1] if "." in name else ""

    if not content:
        return ""

    # Respect explicit engine selection when possible.
    if engine == RecognitionEngine.docling:
        try:
            return (_extract_docling_text(filename, content) or "").strip()
        except Exception:
            return ""

    if ext == "pdf":
        # Fast path: digital PDF text
        try:
            t = (_extract_pdf_text(content) or "").strip()
            if t:
                return t
        except Exception:
            pass

        # Table extraction may catch some PDFs
        if engine in {RecognitionEngine.auto, RecognitionEngine.structured}:
            try:
                t = (_extract_pdf_tables_text(content) or "").strip()
                if t:
                    return t
            except Exception:
                pass

        # OCR via parser_service first
        try:
            t = (_ocr_via_parser_service(filename, content) or "").strip()
            if t:
                return t
        except Exception:
            pass

        # Local OCR fallback
        try:
            t = (_ocr_pdf_bytes(content) or "").strip()
            if t:
                return t
        except Exception:
            pass

        return ""

    if ext in {"png", "jpg", "jpeg"}:
        try:
            t = (_ocr_via_parser_service(filename, content) or "").strip()
            if t:
                return t
        except Exception:
            pass
        try:
            t = (_ocr_image_bytes(content) or "").strip()
            if t:
                return t
        except Exception:
            pass

        if EASYOCR_AVAILABLE:
            try:
                t = (smart_extract_text(filename, content) or "").strip()
                if t:
                    return t
            except Exception:
                pass
        return ""

    if ext == "docx":
        try:
            return (_extract_docx_text(content) or "").strip()
        except Exception:
            return ""

    if ext in {"xlsx", "xls"}:
        try:
            return (_extract_xlsx_text(content) or "").strip()
        except Exception:
            return ""

    # Plain text fallback
    try:
        return content.decode("utf-8", errors="ignore").strip()
    except Exception:
        return ""


def _env_flag(name: str, default: bool = False) -> bool:
    v = (os.getenv(name) or "").strip().lower()
    if not v:
        return default
    return v in {"1", "true", "yes", "on"}


def _groq_cache_enabled() -> bool:
    return _env_flag("GROQ_CACHE_ENABLED", False)


def _groq_cache_dir() -> Path:
    base = (os.getenv("GROQ_CACHE_DIR") or "").strip()
    if base:
        return Path(base)
    return Path(_repo_root) / ".cache" / "groq"


def _groq_cache_ttl_sec() -> int:
    try:
        return int(os.getenv("GROQ_CACHE_TTL_SEC", "86400"))
    except Exception:
        return 86400


def _groq_cache_key(*, kind: str, model: str, system: str, payload_text: str) -> str:
    h = hashlib.sha256()
    h.update((kind or "").encode("utf-8", errors="ignore"))
    h.update(b"\n")
    h.update((model or "").encode("utf-8", errors="ignore"))
    h.update(b"\n")
    h.update((system or "").encode("utf-8", errors="ignore"))
    h.update(b"\n")
    h.update((payload_text or "").encode("utf-8", errors="ignore"))
    return h.hexdigest()


def _groq_cache_read(key: str) -> Optional[dict]:
    if not _groq_cache_enabled():
        return None
    cache_dir = _groq_cache_dir()
    try:
        cache_dir.mkdir(parents=True, exist_ok=True)
    except Exception:
        return None
    path = cache_dir / f"{key}.json"
    try:
        if not path.exists():
            return None
        ttl = _groq_cache_ttl_sec()
        if ttl > 0:
            age = time.time() - path.stat().st_mtime
            if age > ttl:
                return None
        data = json.loads(path.read_text(encoding="utf-8"))
        if not isinstance(data, dict):
            return None
        return data
    except Exception:
        return None


def _groq_cache_write(key: str, data: dict) -> None:
    if not _groq_cache_enabled():
        return
    if not isinstance(data, dict):
        return
    cache_dir = _groq_cache_dir()
    try:
        cache_dir.mkdir(parents=True, exist_ok=True)
    except Exception:
        return
    path = cache_dir / f"{key}.json"
    try:
        payload = json.dumps(data, ensure_ascii=False)
        with tempfile.NamedTemporaryFile("w", delete=False, encoding="utf-8", dir=str(cache_dir), prefix=path.name, suffix=".tmp") as f:
            f.write(payload)
            tmp_name = f.name
        os.replace(tmp_name, str(path))
    except Exception:
        try:
            if "tmp_name" in locals() and tmp_name and os.path.exists(tmp_name):
                os.unlink(tmp_name)
        except Exception:
            pass


def extract_item_names_via_groq_with_usage(*, text: str, api_key: str, max_retries: int = 3) -> tuple[List[str], dict]:
    """Extract item names via Groq with retry on rate limit (429)."""
    import re as _re
    
    raw = (text or "").strip()
    key = (api_key or "").strip()
    if not raw:
        return [], {}
    if not key:
        raise RecognitionDependencyError("GROQ_API_KEY is not configured")

    # Minimize data sent: take only the first chunk - reduced to 8000 for lower TPM
    max_chars = int(os.getenv("GROQ_INPUT_MAX_CHARS", "8000"))
    payload_text = raw[: max(1, min(max_chars, 50000))]

    model = (os.getenv("GROQ_MODEL") or "llama-3.1-8b-instant").strip()
    url = (os.getenv("GROQ_BASE_URL") or "https://api.groq.com/openai/v1").rstrip("/") + "/chat/completions"

    system = (
        "Извлекай ТОЛЬКО базовые наименования позиций (товаров/работ) из текста документа. "
        "Верни ТОЛЬКО JSON без markdown и без пояснений. "
        "Схема строго: {\"names\": [string, ...]}. "
        "Каждый элемент names — КРАТКОЕ наименование без: "
        "  - размеров, диаметров, длин (d=, DN, Ø, мм, м)"
        "  - материалов (ПНД, ПП, ПЭ, Zn)"
        "  - цветов (красная, синяя, черная)"
        "  - технических параметров (SDR, L=, D=, толщина, вес)"
        "  - артикулов, кодов (SN, PN, A60)"
        "  - единиц измерения, цен, количеств"
        "Цифры допускаются ТОЛЬКО если они неотъемлемая часть маркировки (например 'A60', '1000x1000'). "
        "Примеры: "
        "  Правильно: 'Труба жесткая термостойкая', 'Муфта защитная', 'Отвод сварной' "
        "  Неправильно: 'Труба жесткая термостойкая ПНД d=160мм красная', 'Муфта DN50', 'Труба 100мм' "
        "Не добавляй позиции, которых нет в тексте. Извлекай только уникальные наименования."
    )

    user = payload_text

    cache_key = _groq_cache_key(kind="item_names", model=model, system=system, payload_text=payload_text)
    cached = _groq_cache_read(cache_key)
    if isinstance(cached, dict):
        cached_names = cached.get("names")
        cached_usage = cached.get("usage")
        if isinstance(cached_names, list):
            try:
                names = normalize_item_names([str(x or "").strip() for x in cached_names])
            except Exception:
                names = []
            return names, cached_usage if isinstance(cached_usage, dict) else {}

    payload = {
        "model": model,
        "temperature": 0,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "response_format": {"type": "json_object"},
    }

    last_error = None
    usage = {}
    content = ""
    
    # Retry with exponential backoff for rate limit (429)
    for attempt in range(max_retries):
        try:
            with httpx.Client(timeout=30.0) as client:
                r = client.post(url, headers={"Authorization": f"Bearer {key}"}, json=payload)
                
                # Handle rate limit (429)
                if r.status_code == 429:
                    wait_time = 0
                    try:
                        error_data = r.json()
                        error_msg = error_data.get("error", {}).get("message", "")
                        match = _re.search(r"try again in (\d+(?:\.\d+)?)s", error_msg)
                        if match:
                            wait_time = float(match.group(1))
                        else:
                            wait_time = (2 ** attempt) * 2
                    except Exception:
                        wait_time = (2 ** attempt) * 2
                    
                    if attempt < max_retries - 1:
                        logger.info(f"Groq rate limit (429), waiting {wait_time:.1f}s before retry {attempt+1}/{max_retries}")
                        time.sleep(wait_time)
                        continue
                    else:
                        raise RecognitionDependencyError(
                            f"Groq rate limit (429) after {max_retries} retries. Please try again later."
                        )
                
                try:
                    r.raise_for_status()
                except httpx.HTTPStatusError as e:
                    snippet = ""
                    try:
                        snippet = (r.text or "").strip()
                    except Exception:
                        snippet = ""
                    if snippet:
                        snippet = snippet[:400]
                    raise RecognitionDependencyError(
                        f"Groq request failed: {r.status_code} {r.reason_phrase}. {snippet}".strip()
                    ) from e
                    
                data = r.json() or {}
                usage = data.get("usage") or {}
                content = (((data.get("choices") or [{}])[0].get("message") or {}).get("content") or "").strip()
                break
                
        except RecognitionDependencyError:
            last_error = sys.exc_info()[1]
            if attempt < max_retries - 1:
                continue
            raise
        except Exception as e:
            last_error = e
            if attempt < max_retries - 1:
                logger.warning(f"Groq request failed (attempt {attempt+1}/{max_retries}): {e}")
                time.sleep((2 ** attempt) * 1.5)
                continue
            raise RecognitionDependencyError(f"Groq request failed after {max_retries} retries: {e}")

    if not content:
        return [], usage if isinstance(usage, dict) else {}

    # Parse strict JSON object
    obj = None
    try:
        obj = json.loads(content)
    except Exception:
        # Try extracting JSON object from fenced output
        start = content.find("{")
        end = content.rfind("}")
        if start != -1 and end != -1 and end > start:
            try:
                obj = json.loads(content[start : end + 1])
            except Exception:
                obj = None

    if not isinstance(obj, dict):
        return [], usage if isinstance(usage, dict) else {}

    arr = obj.get("names")
    if not isinstance(arr, list):
        return [], usage if isinstance(usage, dict) else {}

    out: List[str] = []
    seen: set[str] = set()
    limit = int(os.getenv("GROQ_OUTPUT_MAX_ITEMS", "200"))
    for it in arr:
        s = str(it or "").strip()
        if not s:
            continue
        s = " ".join(s.split())
        if len(s) < 2:
            continue
        low = s.lower()
        if low in seen:
            continue
        seen.add(low)
        out.append(s)
        if len(out) >= limit:
            break
    final_names = normalize_item_names(out)
    final_usage = usage if isinstance(usage, dict) else {}
    _groq_cache_write(cache_key, {"names": final_names, "usage": final_usage})
    return final_names, final_usage


def extract_search_keys_via_groq(*, text: str, api_key: str, max_retries: int = 5) -> Tuple[List[str], List[str], dict]:
    """
    Extract GROUPED search keys for parsing optimization.
    Returns:
        - search_keys: 3-8 optimized search queries for parsing
        - categories: unique product categories found  
        - usage: token usage info
    """
    import time
    import re as _re

    raw = (text or "").strip()
    key = (api_key or "").strip()
    if not raw:
        return [], [], {}
    if not key:
        raise RecognitionDependencyError("GROQ_API_KEY is not configured")

    max_chars = int(os.getenv("GROQ_INPUT_MAX_CHARS", "8000"))
    payload_text = raw[: max(1, min(max_chars, 50000))]

    model = (os.getenv("GROQ_MODEL") or "llama-3.1-8b-instant").strip()
    url = (os.getenv("GROQ_BASE_URL") or "https://api.groq.com/openai/v1").rstrip("/") + "/chat/completions"

    system = (
        "Ты эксперт по анализу заявок на товары. Твоя задача - сгруппировать позиции и создать МИНИМАЛЬНОЕ количество ПОИСКОВЫХ КЛЮЧЕЙ для поиска поставщиков.\n\n"
        "КРИТИЧНО: Количество ключей должно быть МЕНЬШЕ количества позиций в тексте!\n\n"
        "Формат ответа СТРОГО JSON:\n"
        "{\n"
        '  "search_keys": ["ключ 1", "ключ 2", ...]\n'
        "}\n\n"
        "ПРАВИЛА создания search_keys:\n"
        "1. ОБЪЕДИНЯЙ все похожие товары в ОДИН поисковый ключ\n"
        "2. Ключ = тип товара + материал (если критично важен)\n"
        "3. НЕ включай: размеры, DN, диаметры, количество, цены, артикулы, параметры\n"
        "4. СТРОГИЕ ЛИМИТЫ (НИКОГДА не превышай!):\n"
        "   - 1-3 позиции в тексте → МАКСИМУМ 1 ключ\n"
        "   - 4-10 позиций → МАКСИМУМ 2-3 ключа\n"
        "   - 11-20 позиций → МАКСИМУМ 3-5 ключей\n"
        "   - 21-50 позиций → МАКСИМУМ 5-8 ключей\n"
        "5. Если все позиции одного типа (только настилы / только трубы) → ОДИН ключ!\n\n"
        "ПРИМЕРЫ:\n"
        "- 20 вариантов 'Настил решетчатый сварной Zn' разных размеров → ['Настил решетчатый сварной Zn']\n"
        "- 5 труб ПП разных DN → ['Труба ПП гофрированная']\n"
        "- 3 муфты ПЭ разных DN → ['Муфта защитная ПЭ']\n"
        "- 1 отвод → ['Отвод сварной ПЭ']\n"
        "- 5 труб (3xПП + 2xПНД) → ['Труба ПП', 'Труба ПНД'] или ['Труба гофрированная']\n"
        "- 10 позиций (настилы + ступени + трубы) → ['Настил решетчатый', 'Ступень решетчатая', 'Труба ПП']\n\n"
        "ПОМНИ: Чем меньше ключей - тем лучше! Цель - СОКРАТИТЬ количество запросов к поисковикам."
    )

    payload = {
        "model": model,
        "temperature": 0,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": payload_text},
        ],
        "response_format": {"type": "json_object"},
    }

    last_error = None
    usage = {}
    content = ""

    cache_key = _groq_cache_key(kind="search_keys", model=model, system=system, payload_text=payload_text)
    cached = _groq_cache_read(cache_key)
    if isinstance(cached, dict):
        cached_keys = cached.get("search_keys")
        cached_usage = cached.get("usage")
        if isinstance(cached_keys, list):
            try:
                keys = [str(x or "").strip() for x in cached_keys if str(x or "").strip()]
            except Exception:
                keys = []
            return keys[:8], [], cached_usage if isinstance(cached_usage, dict) else {}

    for attempt in range(max_retries):
        try:
            with httpx.Client(timeout=45.0) as client:
                r = client.post(url, headers={"Authorization": f"Bearer {key}"}, json=payload)

                if r.status_code == 429:
                    wait_time = 3.0
                    try:
                        error_data = r.json()
                        error_msg = error_data.get("error", {}).get("message", "")
                        match = _re.search(r"try again in (\d+)(?:ms|s)", error_msg, _re.IGNORECASE)
                        if match:
                            wait_val = int(match.group(1))
                            if "ms" in error_msg.lower():
                                wait_time = wait_val / 1000.0 + 1.0
                            else:
                                wait_time = wait_val + 1.0
                        else:
                            wait_time = (attempt + 1) * 3.0
                    except Exception:
                        wait_time = (attempt + 1) * 3.0

                    logger.warning(f"Groq rate limit hit, waiting {wait_time:.1f}s before retry {attempt + 1}/{max_retries}")
                    time.sleep(wait_time)
                    continue

                try:
                    r.raise_for_status()
                except httpx.HTTPStatusError as e:
                    snippet = ""
                    try:
                        snippet = (r.text or "").strip()[:400]
                    except Exception:
                        pass
                    last_error = RecognitionDependencyError(
                        f"Groq request failed: {r.status_code} {r.reason_phrase}. {snippet}".strip()
                    )
                    if attempt < max_retries - 1:
                        time.sleep((attempt + 1) * 1.0)
                        continue
                    raise last_error from e

                data = r.json() or {}
                usage = data.get("usage") or {}
                content = (((data.get("choices") or [{}])[0].get("message") or {}).get("content") or "").strip()
                break

        except RecognitionDependencyError:
            raise
        except Exception as e:
            last_error = RecognitionDependencyError(f"Groq request failed: {e}")
            if attempt < max_retries - 1:
                time.sleep((attempt + 1) * 1.0)
                continue
            raise last_error

    if not content:
        return [], [], usage if isinstance(usage, dict) else {}

    obj = None
    try:
        obj = json.loads(content)
    except Exception:
        start = content.find("{")
        end = content.rfind("}")
        if start != -1 and end != -1 and end > start:
            try:
                obj = json.loads(content[start : end + 1])
            except Exception:
                obj = None

    if not isinstance(obj, dict):
        return [], [], usage if isinstance(usage, dict) else {}

    search_keys = obj.get("search_keys", [])

    if not isinstance(search_keys, list):
        search_keys = []

    keys = []
    seen = set()
    for k in search_keys:
        s = str(k or "").strip()
        if not s or len(s) < 2:
            continue
        s = " ".join(s.split())
        low = s.lower()
        if low in seen:
            continue
        seen.add(low)
        keys.append(s)
        if len(keys) >= 8:
            break

    final_usage = usage if isinstance(usage, dict) else {}
    _groq_cache_write(cache_key, {"search_keys": keys, "usage": final_usage})
    return keys, [], final_usage


def parse_positions_from_text(text: str) -> List[str]:
    import re

    raw = text or ""
    lines = [re.sub(r"\s+", " ", (ln or "").strip()) for ln in re.split(r"\r?\n", raw)]
    lines = [ln for ln in lines if ln]

    # Normalize common OCR/table artifacts:
    # - "1Настил ..." -> "1 Настил ..."
    # - "21. Настил ..." / "21) Настил ..." -> "21 Настил ..." (handled after splitting glued rows)
    lines = [re.sub(r"^(\d{1,3})(?=\D)", r"\1 ", ln) for ln in lines]
    # - "2/ Настил ..." -> "2 Настил ..."
    # - "0 |Настил ..." -> "0 Настил ..."
    lines = [re.sub(r"^(\d{1,3})\s*[/|\\]\\s*", r"\1 ", ln) for ln in lines]

    header_markers = [
        r"наименование\s+продукц",
        r"кол-?во",
        r"складе",
        r"вес\s+издел",
        r"цена\s+за\s+ед",
        r"цена\s+за\s+м\s*2",
        r"цена\s+за\s+тонн",
        r"руб\b",
        r"ндс",
    ]

    def _is_header_like(s: str) -> bool:
        ls = s.lower()
        if re.search(r"^(итого|всего\s+к\s+оплате|сумма\s+ндс|всего\s+наименований)\b", ls):
            return True
        return any(re.search(p, ls) for p in header_markers)

    # Heuristic: a line-item usually contains either:
    # - product keywords (Настил/Ступень/Профиль/Решетчатый), OR
    # - a dimension pattern like 34x38/30x3, 1000х1000, 33*33/30*2, OR
    # - the phrase "по запросу" (common in price lists)
    product_kw = re.compile(
        r"\b(настил|ступен\w*|профил\w*|решетчат\w*|реш[её]тк\w*|муфт\w*|труб\w*|отвод\w*|заглушк\w*|переход\w*|тройник\w*|болт\w*|гайк\w*|шайб\w*|позиц\w*)\b",
        re.IGNORECASE,
    )
    dim_kw = re.compile(r"\b\d{1,4}\s*[xх\*]\s*\d{1,4}", re.IGNORECASE)
    has_money = re.compile(r"\b\d{1,3}(?:[\s\u00A0]\d{3})*(?:[\.,]\d{2})?\b")

    def _looks_like_position_line(s: str) -> bool:
        if not s or len(s) < 8:
            return False
        if _is_header_like(s):
            return False
        if re.match(r"^\s*\d{1,3}\s+\D", s) and re.search(r"[A-Za-zА-Яа-я]", s):
            return True
        ls = s.lower()
        s_wo_ordinal = re.sub(r"^\s*\d{1,3}\s*[\)\.]?\s*", "", s)
        if (
            product_kw.search(s)
            and not has_money.search(s_wo_ordinal)
            and re.search(r"\b(изготовлен\w*|производ\w*|соответств\w*|примечан\w*)\b", ls)
        ):
            return False
        if product_kw.search(s) or dim_kw.search(s) or ("по запросу" in s.lower()):
            return True
        # fallback: if it has several numeric tokens, it may be a table row
        nums = has_money.findall(s)
        return len(nums) >= 3

    def _starts_like_new_item(s: str) -> bool:
        # Many OCR engines drop the leftmost ordinal column. For table rows we still
        # want to split items by product row starts.
        ls = (s or "").strip().lower()
        if not ls or _is_header_like(ls):
            return False
        if re.match(r"^\d{1,3}\s+", s):
            return True
        if re.match(r"^\d{1,3}\s*[/|\\]\\s*", s):
            return True
        if re.match(r"^(настил|ступен\w*|профил\w*|решетчат\w*|реш[её]тк\w*|муфт\w*|труб\w*|отвод\w*|заглушк\w*|переход\w*|тройник\w*)\b", ls):
            return _looks_like_position_line(s)
        return False

    # If a line contains multiple '|' segments (table rows glued), split into chunks.
    expanded: List[str] = []
    for ln in lines:
        # Some OCR outputs glue multiple numbered rows into a single long line.
        # Split by embedded patterns like " 21. " or " 21) ".
        if re.search(r"(?:^|\s)\d{1,3}[\).]\s+", ln):
            chunks = [c.strip() for c in re.split(r"(?=(?:^|\s)\d{1,3}[\).]\s+)", ln) if (c or "").strip()]
            for c in chunks:
                expanded.append(re.sub(r"^(\d{1,3})[\).]\s*", r"\1 ", c))
            continue
        if "|" in ln:
            parts = [p.strip() for p in ln.split("|") if p.strip()]
            for p in parts:
                expanded.append(p)
        else:
            expanded.append(ln)

    # Normalize dot/paren ordinal prefixes after splitting
    expanded = [re.sub(r"^(\d{1,3})[\).]\s*", r"\1 ", ln) for ln in expanded]

    # Remove obvious header-like lines early
    expanded = [ln for ln in expanded if not _is_header_like(ln)]

    item_re = re.compile(r"^(\d{1,3})\s+.+")
    has_items = any(_starts_like_new_item(ln) for ln in expanded)
    if has_items:
        out: List[str] = []
        current: Optional[str] = None
        for ln in expanded:
            # Skip garbage fragments like "000,00]" or empty bracket-only tokens
            if re.fullmatch(r"[\[\]\)\(\d\s,\.]+", ln):
                continue

            # Skip pure small numbers like "312"
            if re.fullmatch(r"\d{1,4}", ln):
                continue

            # Drop very short punctuation-only lines like ")".
            if len(ln) <= 3 and re.fullmatch(r"[\)\(\]\[]+", ln):
                continue

            if _starts_like_new_item(ln) and _looks_like_position_line(ln):
                if current:
                    out.append(current)
                current = ln
                continue

            if not current and _looks_like_position_line(ln):
                current = ln
                continue

            if current:
                if _is_header_like(ln):
                    continue
                # Glue short continuations (OCR wrap)
                if len(ln) <= 120:
                    current = f"{current} {ln}".strip()

        if current:
            out.append(current)

        # Final filter to avoid header fragments becoming items
        out = [x for x in out if _looks_like_position_line(x)]
        return out

    # If no explicit numbering, keep only lines that look like positions
    filtered = [ln for ln in expanded if _looks_like_position_line(ln)]
    return filtered if filtered else expanded


def extract_parsing_keys_from_positions(items: List[str]) -> List[str]:
    import re

    def _class_only_enabled() -> bool:
        return (os.getenv("KEY_CLASS_ONLY", "").strip().lower() in {"1", "true", "yes", "on"})

    class _LLMKeyExtraction(BaseModel):
        product: str = Field(default="")
        attrs: List[str] = Field(default_factory=list)
        codes: List[str] = Field(default_factory=list)
        variants: List[str] = Field(default_factory=list)

    def _extract_json_object(s: str) -> Optional[str]:
        if not s:
            return None
        t = s.strip()
        if "```" in t:
            start = t.find("{")
            end = t.rfind("}")
            if start != -1 and end != -1 and end > start:
                t = t[start : end + 1]
        start = t.find("{")
        end = t.rfind("}")
        if start == -1 or end == -1 or end <= start:
            return None
        return t[start : end + 1]

    def _ollama_extract(line: str) -> Optional[_LLMKeyExtraction]:
        if not settings.llm_keys_enabled:
            return None
        if not settings.ollama_model:
            return None
        url = (settings.ollama_url or "").rstrip("/") + "/api/generate"
        if _class_only_enabled():
            system = (
                "Ты извлекаешь ключ-класс номенклатуры из одной строки OCR. "
                "Верни ТОЛЬКО JSON-объект без текста/пояснений/markdown. "
                "Схема: {product: string, attrs: string[], codes: string[], variants: string[]}. "
                "product должен быть кратким названием класса (например: 'Муфта защитная', 'Труба ПП гофрированная с раструбом'). "
                "Не включай DN/ID/SDR/числа/размеры/длины/мм/м/шт в product/attrs/codes/variants. "
                "attrs: только значимые признаки (например: гофрированная, с раструбом, защитная). "
                "codes: только коды типа SP, S4, S11. "
                "variants: СО/ОСО если есть."
            )
        else:
            system = (
                "Ты экстрактор ключей номенклатуры из одной строки OCR. "
                "Верни ТОЛЬКО JSON-объект. Никакого текста/пояснений/markdown. "
                "Схема: {product: string, attrs: string[], codes: string[], variants: string[]}. "
                "Не включай размеры/габариты/числа как атрибуты. "
                "В attrs добавляй только значимые признаки (например: решетчатый, сварной, защитная, электросварной). "
                "В codes добавляй коды типа SP, S4, S11. "
                "В variants добавляй СО/ОСО если есть."
            )
        prompt = f"Строка: {line}"

        payload = {
            "model": settings.ollama_model,
            "prompt": prompt,
            "system": system,
            "stream": False,
            "format": "json",
            "options": {"temperature": 0},
        }

        def _ollama_extract_cli() -> Optional[_LLMKeyExtraction]:
            try:
                full_prompt = system + "\n" + prompt
                completed = subprocess.run(
                    ["ollama", "run", settings.ollama_model, full_prompt],
                    capture_output=True,
                    text=True,
                    encoding="utf-8",
                    errors="replace",
                    timeout=float(settings.ollama_timeout_sec),
                )
            except Exception:
                return None
            raw_cli = (completed.stdout or "").strip()
            js_cli = _extract_json_object(raw_cli)
            if not js_cli:
                return None
            try:
                obj = _LLMKeyExtraction.model_validate_json(js_cli)
                logger.info("ollama_extract transport=cli ok=1")
                return obj
            except ValidationError:
                return None

        try:
            with httpx.Client(timeout=float(settings.ollama_timeout_sec)) as client:
                r = client.post(url, json=payload)
                r.raise_for_status()
                data = r.json()
        except Exception:
            return _ollama_extract_cli()

        raw = (data or {}).get("response")
        js = _extract_json_object(raw or "")
        if not js:
            return _ollama_extract_cli()
        try:
            obj = _LLMKeyExtraction.model_validate_json(js)
            logger.info("ollama_extract transport=http ok=1")
            return obj
        except ValidationError:
            return _ollama_extract_cli()

    def _normalize_llm_attr(a: str) -> Optional[str]:
        if not a:
            return None
        t = _cyrillize_similar(a).strip()
        if not t:
            return None
        low = t.lower()
        if re.search(r"\bzn\b", low):
            return "Zn"
        if re.search(r"\bоцинк\w*\b", low):
            return "оцинк"
        if re.search(r"\bпо\s+запросу\b", low):
            return "по запросу"
        if re.search(r"\bобр\.?\b", low) or re.search(r"\bобрам\w*\b", low):
            return "обр"
        m = re.search(r"\bтип\s*([aа])\b", low)
        if m:
            return "тип А"
        return None

    def _line_has_token(src_line: str, token: str) -> bool:
        if not src_line or not token:
            return False
        l = _latinize_similar(src_line)
        c = _cyrillize_similar(src_line)
        t_lat = _latinize_similar(token)
        t_cyr = _cyrillize_similar(token)
        return bool(
            re.search(rf"\b{re.escape(t_lat)}\b", l, flags=re.IGNORECASE)
            or re.search(rf"\b{re.escape(t_cyr)}\b", c, flags=re.IGNORECASE)
        )

    def _latinize_similar(s: str) -> str:
        table = str.maketrans(
            {
                "А": "A",
                "В": "B",
                "Е": "E",
                "К": "K",
                "М": "M",
                "Н": "H",
                "О": "O",
                "Р": "P",
                "С": "C",
                "Т": "T",
                "Х": "X",
                "а": "a",
                "в": "b",
                "е": "e",
                "к": "k",
                "м": "m",
                "н": "h",
                "о": "o",
                "р": "p",
                "с": "c",
                "т": "t",
                "х": "x",
            }
        )
        return (s or "").translate(table)

    def _cyrillize_similar(s: str) -> str:
        # Convert latin lookalikes to cyrillic to stabilize Russian word detection
        table = str.maketrans(
            {
                "A": "А",
                "B": "В",
                "C": "С",
                "E": "Е",
                "H": "Н",
                "K": "К",
                "M": "М",
                "O": "О",
                "P": "Р",
                "T": "Т",
                "X": "Х",
                "a": "а",
                "b": "в",
                "c": "с",
                "e": "е",
                "h": "н",
                "k": "к",
                "m": "м",
                "o": "о",
                "p": "р",
                "t": "т",
                "x": "х",
            }
        )
        return (s or "").translate(table)

    def _clean_prefix(line: str) -> str:
        s = (line or "").strip()
        s = re.sub(
            r"^\s*(позиции|добавить\s+позици\w+|файл\s+заявки\s+для\s+распознавания)\b\s*",
            "",
            s,
            flags=re.IGNORECASE,
        )
        s = re.sub(r"^\s*\d{1,3}\s*([/|\\\.)-]?\s*)", "", s)
        s = re.sub(r"\s+", " ", s).strip()
        return s

    def _tokens(line: str) -> List[str]:
        parts = re.findall(r"[A-Za-zА-Яа-я0-9]+", line or "")
        return [p for p in (p.strip() for p in parts) if p]

    def _canonical_product(word: str) -> Optional[str]:
        w = _cyrillize_similar((word or "").strip()).lower()
        if w.startswith("настил"):
            return "Настил"
        if w.startswith("ступ"):
            return "Ступень"
        if w.startswith("проф"):
            return "Профиль"
        if w.startswith("решетк"):
            return "Решетка"
        if w.startswith("флан"):
            return "Фланец"
        if w.startswith("муфт"):
            return "Муфта"
        if w.startswith("отвод"):
            return "Отвод"
        if w.startswith("заглуш"):
            return "Заглушка"
        if w.startswith("переход"):
            return "Переход"
        if w.startswith("тройн") or w.startswith("троин"):
            return "Тройник"
        if w.startswith("кран"):
            return "Кран"
        if w.startswith("труб"):
            return "Труба"
        return None

    def _code_token(tok: str) -> Optional[str]:
        t = _latinize_similar((tok or "").strip()).upper()
        if not t:
            return None
        m = re.fullmatch(r"DN\s*(\d{1,4})", t)
        if m:
            return f"DN{m.group(1)}"
        m = re.fullmatch(r"ID\s*(\d{1,4})", t)
        if m:
            return f"ID{m.group(1)}"
        m = re.fullmatch(r"SDR\s*(\d{1,3})", t)
        if m:
            return f"SDR{m.group(1)}"
        if re.fullmatch(r"S\d{1,3}", t):
            return t
        if re.fullmatch(r"SP", t):
            return "SP"
        if re.fullmatch(r"OSO", t) or re.fullmatch(r"ОСО", tok.strip().upper()):
            return "ОСО"
        return None

    def _variant_token(tok: str) -> Optional[str]:
        # ОСО/СО can be OCR'ed as latin OCO/CO. Normalize via cyrillize.
        t = _cyrillize_similar((tok or "").strip()).upper()
        if t in {"ОСО", "СО"}:
            return t
        return None

    def _controller_repair_key(pretty: str, src_line: str) -> str:
        if not pretty:
            return pretty

        key = pretty
        l = _latinize_similar(src_line or "")
        c = _cyrillize_similar(src_line or "")
        if re.search(r"\bS4\b", l, flags=re.IGNORECASE) and not re.search(r"\bS4\b", key, flags=re.IGNORECASE):
            key = f"{key} S4"
        if re.search(r"\bSP\b", l, flags=re.IGNORECASE) and not re.search(r"\bSP\b", key, flags=re.IGNORECASE):
            key = f"{key} SP"
        if re.search(r"\bПЭ\b", c, flags=re.IGNORECASE) and "ПЭ" not in key:
            key = f"{key} ПЭ"
        if re.search(r"\bПП\b", c, flags=re.IGNORECASE) and "ПП" not in key:
            key = f"{key} ПП"
        if re.search(r"\bОСО\b", c, flags=re.IGNORECASE) and "ОСО" not in key:
            key = f"{key} ОСО"
        if re.search(r"\bСО\b", c, flags=re.IGNORECASE) and "СО" not in key:
            key = f"{key} СО"

        if _class_only_enabled():
            key = re.sub(r"\b(DN\s*\d{1,4}|ID\s*\d{1,4}|SDR\s*\d{1,3})\b", " ", key, flags=re.IGNORECASE)
            key = re.sub(r"\s+", " ", key).strip()
            return key

        dn = None
        m = re.search(r"\bDN\s*(\d{1,4})\b", l, flags=re.IGNORECASE)
        if m:
            dn = f"DN{m.group(1)}"
        if dn and not re.search(rf"\b{re.escape(dn)}\b", key, flags=re.IGNORECASE):
            key = f"{key} {dn}"

        ident = None
        m = re.search(r"\bID\s*(\d{1,4})\b", l, flags=re.IGNORECASE)
        if m:
            ident = f"ID{m.group(1)}"
        if ident and not re.search(rf"\b{re.escape(ident)}\b", key, flags=re.IGNORECASE):
            key = f"{key} {ident}"

        key = re.sub(r"\s+", " ", key).strip()
        return key

    def _looks_like_dimension(tok: str) -> bool:
        if not tok:
            return False
        t = _latinize_similar(tok)
        # Keep product codes like S4/SP/S11 as non-dimensions (processed separately)
        if re.fullmatch(r"(?:SP|S\d{1,3})", t, flags=re.IGNORECASE):
            return False
        # Pure numbers are usually dimensions/quantities/etc.
        if re.fullmatch(r"\d{1,6}", t):
            return True
        # Any explicit dimension separators
        if re.search(r"[xх\*/]", tok, flags=re.IGNORECASE):
            return True
        return False

    def _normalize_adj(word: str) -> Optional[str]:
        w = _cyrillize_similar((word or "").strip()).lower()
        if not w:
            return None
        if w.startswith("решетчат"):
            return "решетчатый"
        if w.startswith("сварн"):
            return "сварной"
        if w.startswith("защитн"):
            return "защитная"
        if w.startswith("электросварн"):
            return "электросварной"
        return None

    def _normalize_material(word: str) -> Optional[str]:
        w = _cyrillize_similar((word or "").strip()).lower()
        if not w:
            return None
        if w in {"пэ", "пе"}:
            return "ПЭ"
        if w == "пп":
            return "ПП"
        if w.startswith("медн"):
            return "медный"
        if w.startswith("бронз"):
            return "бронзовый"
        if w.startswith("алюм"):
            return "алюминиевый"
        if w.startswith("нерж") or w.startswith("нержав"):
            return "нержавеющий"
        return None

    def _looks_like_adjective(word: str) -> bool:
        w = _cyrillize_similar((word or "").strip()).lower()
        if not w or not re.search(r"[а-яё]", w):
            return False
        return bool(
            re.search(
                r"(ый|ий|ой|ая|яя|ое|ее|ые|ие|ого|его|ому|ему|ым|им|ую|юю|ой|ей)$",
                w,
            )
        )

    stop_prefixes = {
        "гост",
        "ту",
        "запросу",
    }

    stop_tokens = {
        "для",
        "через",
        "под",
        "над",
        "при",
        "из",
        "на",
        "в",
        "к",
        "от",
        "по",
        "мм",
        "шт",
        "шт.",
        "кг",
        "г",
        "м",
        "м2",
        "м3",
        "ж",
        "б",
        "dn",
        "id",
        "sdr",
        "sdr17",
        "sdr11",
        "d",
    }

    def _is_bad_product_token(tok: str) -> bool:
        t = _cyrillize_similar((tok or "").strip()).lower()
        if not t:
            return True
        if t in stop_tokens:
            return True
        if any(t.startswith(p) for p in stop_prefixes):
            return True
        if re.fullmatch(r"[a-zа-я]{1,3}", t):
            return True
        return False

    out: List[str] = []
    seen: set[str] = set()

    best_by_product: dict[str, int] = {}

    for raw in items or []:
        line = _clean_prefix(raw)
        if not line:
            continue

        ls = line.lower()
        if re.search(r"\b(изготовлен\w*|производ\w*|соответств\w*|примечан\w*)\b", ls):
            nums = re.findall(r"\b\d{1,3}(?:[\s\u00A0]\d{3})*(?:[\.,]\d{2})?\b", line)
            if len(nums) < 2:
                continue
        toks = _tokens(line)
        if not toks:
            continue

        prod: Optional[str] = None
        prod_idx: Optional[int] = None
        for i, tok in enumerate(toks[:14]):
            if _looks_like_dimension(tok):
                continue
            low = _cyrillize_similar(tok).lower()
            if low in stop_tokens:
                continue
            if any(low.startswith(p) for p in stop_prefixes):
                continue
            p = _canonical_product(tok)
            if p:
                prod = p
                prod_idx = i
                break

        if not prod:
            for i, tok in enumerate(toks[:10]):
                if _looks_like_dimension(tok):
                    continue
                if _is_bad_product_token(tok):
                    continue
                if _normalize_adj(tok) or _normalize_material(tok) or _looks_like_adjective(tok):
                    continue
                if re.fullmatch(r"[A-Za-zА-Яа-я]{2}", tok or ""):
                    continue
                prod = _cyrillize_similar(tok).capitalize()
                prod_idx = i
                break

        if not prod:
            continue

        attrs: List[str] = []
        codes: List[str] = []
        variants: List[str] = []

        scan_from = (prod_idx + 1) if prod_idx is not None else 0
        scan_to = min(scan_from + 60, len(toks))
        for tok in toks[scan_from:scan_to]:
            if not tok:
                continue
            if _looks_like_dimension(tok):
                continue

            mat = _normalize_material(tok)
            if mat and mat not in attrs:
                attrs.append(mat)
                continue

            adj = _normalize_adj(tok)
            if adj and adj not in attrs:
                attrs.append(adj)
                continue

            v = _variant_token(tok)
            if v and v not in variants:
                variants.append(v)
                continue

            c = _code_token(tok)
            if c and c not in codes:
                codes.append(c)
                continue

        if not any(a in attrs for a in {"ПЭ", "ПП"}):
            for tok in toks:
                mat = _normalize_material(tok)
                if mat in {"ПЭ", "ПП"} and mat not in attrs:
                    attrs.append(mat)

        if "сварной" in attrs and "решетчатый" in attrs:
            attrs = [a for a in attrs if a not in {"решетчатый", "сварной"}] + ["решетчатый", "сварной"]

        key_parts: List[str] = [prod]
        key_parts.extend(attrs)
        key_parts.extend(codes)
        key_parts.extend(variants)

        pretty = " ".join(key_parts).strip()
        pretty = re.sub(r"\s+", " ", pretty)
        pretty = _controller_repair_key(pretty, line)
        if not pretty:
            continue

        if _class_only_enabled():
            pretty = re.sub(r"\b(DN\s*\d{1,4}|ID\s*\d{1,4}|SDR\s*\d{1,3})\b", " ", pretty, flags=re.IGNORECASE)
            pretty = re.sub(r"\s+", " ", pretty).strip()

        if settings.llm_keys_enabled and settings.ollama_model:
            low_line = _cyrillize_similar(line).lower()
            looks_like_rich = bool(
                re.search(r"\b(пэ|пп|осо|со|s\d{1,3}|sp)\b", low_line)
                or re.search(r"\b(решетчат|реш[её]тк|сварн|защитн|электросварн)\b", low_line)
                or re.search(r"\b(zn|оцинк|по\s+запросу|обр\.?|обрам)\b", low_line)
                or re.search(r"\b(настил|ступен\w*)\s+[рp]\b", low_line)
            )
            salient_in_line = bool(re.search(r"\b(dn\d{1,4}|sdr\d{1,3}|s\d{1,3}|sp)\b", low_line))
            salient_in_key = bool(re.search(r"\b(dn\d{1,4}|sdr\d{1,3}|s\d{1,3}|sp)\b", pretty.lower()))
            needs_enrichment = (len(pretty.split()) <= 1 or (len(attrs) + len(codes) + len(variants) == 0)) or (
                salient_in_line and not salient_in_key
            )
            force_llm = bool(getattr(settings, "llm_keys_force", False))
            if (looks_like_rich and needs_enrichment) or force_llm:
                llm = _ollama_extract(line)
                if llm and (llm.product or llm.attrs or llm.codes or llm.variants):
                    llm_prod = llm.product.strip()
                    if llm_prod:
                        can = _canonical_product(llm_prod) or llm_prod[:1].upper() + llm_prod[1:]
                    else:
                        can = prod
                    llm_attrs = []
                    for a in llm.attrs:
                        na = _normalize_adj(a) or _normalize_material(a) or _normalize_llm_attr(a) or None
                        if na and na not in llm_attrs:
                            llm_attrs.append(na)
                    llm_codes = []
                    for c in llm.codes:
                        nc = _code_token(c)
                        if nc and _line_has_token(line, nc) and nc not in llm_codes:
                            llm_codes.append(nc)
                    llm_vars = []
                    for v in llm.variants:
                        nv = _variant_token(v)
                        if nv and _line_has_token(line, nv) and nv not in llm_vars:
                            llm_vars.append(nv)

                    repaired = " ".join([can] + llm_attrs + llm_codes + llm_vars).strip()
                    repaired = re.sub(r"\s+", " ", repaired)
                    repaired = _controller_repair_key(repaired, line)
                    if repaired:
                        pretty = repaired

        specificity = (len(attrs) * 10) + (len(codes) * 5) + (len(variants) * 3)
        best = best_by_product.get(prod)
        if best is not None and specificity < best and specificity <= 0:
            continue
        best_by_product[prod] = max(best_by_product.get(prod, 0), specificity)

        norm = pretty.lower()
        if norm in seen:
            continue
        seen.add(norm)
        out.append(pretty)

    filtered: List[str] = []
    for k in out:
        base = (k.split(" ", 1)[0] if k else "")
        if not base:
            continue
        if k == base and best_by_product.get(base, 0) > 0:
            continue
        filtered.append(k)

    return filtered


def extract_parsing_keys_per_position(items: List[str]) -> List[str]:
    out: List[str] = []
    for it in items or []:
        it_s = (it or "").strip()
        if not it_s:
            continue
        keys = extract_parsing_keys_from_positions([it_s])
        if keys:
            out.append(keys[0])
        else:
            out.append(it_s)
    return out


def normalize_positions(items: List[str]) -> List[str]:
    import re

    product_kw = re.compile(r"\b(настил|ступен\w*|профил\w*|решетчат\w*|реш[её]тк\w*|позиц\w*)\b", re.IGNORECASE)
    dim_kw = re.compile(r"\b\d{1,4}[\s\u00A0]*[xх\*][\s\u00A0]*\d{1,4}", re.IGNORECASE)
    number_tokens = re.compile(r"\b\d{1,3}(?:[\s\u00A0]\d{3})*(?:[\.,]\d{2})?\b")

    def _truncate_at_stop_phrases(s: str) -> str:
        # Stop phrases that often start footers/totals and must not be part of a line-item,
        # even if OCR/table extraction glued them into the same line.
        stop_patterns = [
            r"\bитого\b\s*:",
            r"\bсумма\s+ндс\b\s*:",
            r"\bвсего\s+к\s+оплате\b\s*:",
            r"\bчетыре\s+миллиона\b",  # spelled-out total (example from user PDF)
            r"\bуслови[яе]\s+поставк[иы]\b",
            r"\bусловие\s+оплаты\b\s*:",
            r"\bсрок\s+производства\b\s*:",
            r"\bстоимость\s+материала\b",
            r"\bконтактн(?:ая|ые)\s+информаци[яи]\b",
            r"\bруководитель\b",
            r"\bглавн(?:ый|ая)\b\s*\(?(?:старший)?\)?\s*бухгалтер\b",
            r"\bответственный\b",
            r"\bдолжность\b",
            r"\bподпись\b",
            r"\bрасшифровк\w+\s+подписи\b",
            r"\bмоб\.?\b",
            r"\bwhatsap\b",
            r"\be-?mail\b",
            r"\bтел\.?\b",
            r"\bhttp\b",
            r"\bwww\.",
            r"\bпредложени\w+\b",
            r"\bгост\b",
            r"\bсто\b",
            r"\bсоответств\w+\b",
            r"\bпокрыти\w+\b",
            r"\bцинк\w+\b\s+по\b",
            r"\bмарка\s+стал\w+\b",
        ]
        low = s.lower()
        cut_at: Optional[int] = None
        for pat in stop_patterns:
            m = re.search(pat, low, flags=re.IGNORECASE)
            if not m:
                continue
            idx = m.start()
            if cut_at is None or idx < cut_at:
                cut_at = idx
        if cut_at is None:
            return s
        return (s[:cut_at] or "").strip()

    out: List[str] = []
    seen: set[str] = set()

    def _short_product_name(s: str) -> str:
        low = (s or "").lower()

        low = re.sub(r"\bреш[её]тчат", "решетчат", low)

        if "настил" in low:
            base = "Настил"
        elif re.search(r"\bступен\w*\b", low):
            base = "Ступень"
        else:
            return (s or "").strip()

        # Only collapse to canonical short name when the line is clearly truncated/too generic.
        # For real table rows keep details (SP, dimensions, prices, etc.).
        has_digit = re.search(r"\d", low) is not None
        has_sp = re.search(r"\bsp\b", low) is not None
        has_dim = re.search(r"\b\d{1,4}\s*[xх\*]\s*\d{1,4}", low) is not None
        is_very_short = len((s or "").strip()) <= 18
        is_generic_only = re.fullmatch(r"\s*(настил|ступен\w*)\s*", low) is not None
        should_canonicalize = is_generic_only or (is_very_short and not (has_digit or has_sp or has_dim))

        if not should_canonicalize:
            return (s or "").strip()

        parts: List[str] = [base]
        # Canonical short name for truncated/generic lines.
        parts.append("решетчатая" if base == "Ступень" else "решетчатый")
        parts.append("сварная" if base == "Ступень" else "сварной")

        return " ".join(parts)
    for s in items or []:
        raw = re.sub(r"[\s\u00A0]+", " ", (s or "").strip())
        if not raw:
            continue

        raw = re.sub(r"^\d{1,3}\s+", "", raw).strip()

        # Aggressively split embedded ordinals that OCR often keeps inside the same line.
        # Example: "... 265 405,14 2. 2 ..." -> separate chunks.
        raw = re.sub(r"(?:^|[\s\u00A0])(\d{1,3})\.(?=[\s\u00A0]+)", r"\n\1 ", raw)
        raw = re.sub(r"(?:^|[\s\u00A0])(\d{1,3})\)(?=[\s\u00A0]+)", r"\n\1 ", raw)
        raw = raw.replace("\r", "")
        raw_lines = [x.strip() for x in raw.split("\n") if x.strip()]

        for raw_line in raw_lines:
            # Split lines where multiple positions got glued together, e.g. "... 20. ... 21. ..."
            # Support both "20." and "20)" formats.
            chunks = [c.strip() for c in re.split(r"(?=(?:^|\s)\d{1,3}[\).]\s+)", raw_line) if (c or "").strip()]

            for chunk in chunks:
                v = (chunk or "").strip()
                if not v:
                    continue

                # Normalize whitespace early (including NBSP)
                v = re.sub(r"[\s\u00A0]+", " ", v).strip()

                # First: truncate glued footers/totals inside the same line.
                v = _truncate_at_stop_phrases(v)
                if not v:
                    continue

                # Remove leading ordinal numbering (e.g. "22. ..." or "22) ...")
                v = re.sub(r"^\d{1,3}[\)\.]\s*", "", v).strip()
                v = re.sub(r"^\d{1,3}\s+", "", v).strip()

                # Remove/split leftover embedded ordinals like " 21. " that survived.
                v = re.sub(r"[\s\u00A0]+\d{1,3}[\).][\s\u00A0]+", " ", v)

                # Drop if it degraded to a pure number
                if re.fullmatch(r"\d{1,6}", v or ""):
                    continue

                # Must contain letters (avoid stray numeric/markup artifacts)
                if not re.search(r"[A-Za-zА-Яа-я]", v):
                    continue

                low = v.lower()

                # If it doesn't look like a product row, require several numeric tokens
                if not (product_kw.search(v) or dim_kw.search(v) or ("по запросу" in low)):
                    nums = number_tokens.findall(v)
                    if len(nums) < 3:
                        continue

                # Hard-drop if it looks like a table header fragment.
                if any(k in low for k in ["наименование", "кол-во", "вес изделия", "цена за", "руб", "ндс"]):
                    continue

                # Remove trailing punctuation-only leftovers.
                v = re.sub(r"[\]\)]+$", "", v).strip()

                # Final collapse
                v = re.sub(r"[\s\u00A0]+", " ", v).strip()

                # Reduce long technical tails to a short product name when possible
                v = _short_product_name(v)

                # Drop pure punctuation remnants
                if not v or re.fullmatch(r"[\)\(\]\[]+", v):
                    continue

                if v in seen:
                    continue
                seen.add(v)
                out.append(v)
    return out


def _decode_text(data: bytes) -> str:
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return str(data)


def _extract_html_text(data: bytes) -> str:
    import re

    raw = _decode_text(data)
    raw = re.sub(r"<script[\s\S]*?</script>", " ", raw, flags=re.IGNORECASE)
    raw = re.sub(r"<style[\s\S]*?</style>", " ", raw, flags=re.IGNORECASE)
    raw = re.sub(r"<[^>]+>", " ", raw)
    raw = re.sub(r"&nbsp;", " ", raw, flags=re.IGNORECASE)
    raw = re.sub(r"&amp;", "&", raw, flags=re.IGNORECASE)
    raw = re.sub(r"&lt;", "<", raw, flags=re.IGNORECASE)
    raw = re.sub(r"&gt;", ">", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s+", " ", raw).strip()
    return raw


def _extract_eml_text(data: bytes) -> str:
    from email import policy
    from email.parser import BytesParser

    msg = BytesParser(policy=policy.default).parsebytes(data or b"")
    parts: List[str] = []
    subj = (msg.get("subject") or "").strip()
    if subj:
        parts.append(subj)
    if msg.is_multipart():
        for p in msg.walk():
            ctype = (p.get_content_type() or "").lower()
            if ctype not in {"text/plain", "text/html"}:
                continue
            try:
                payload = p.get_content()
            except Exception:
                payload = None
            if not payload:
                continue
            if isinstance(payload, bytes):
                t = _decode_text(payload)
            else:
                t = str(payload)
            if ctype == "text/html":
                t = _extract_html_text(t.encode("utf-8", errors="ignore"))
            t = (t or "").strip()
            if t:
                parts.append(t)
    else:
        try:
            payload = msg.get_content()
        except Exception:
            payload = None
        if payload:
            if isinstance(payload, bytes):
                parts.append(_decode_text(payload))
            else:
                parts.append(str(payload))
    return "\n".join([p for p in parts if (p or "").strip()])


def _extract_rtf_text(data: bytes) -> str:
    import re

    raw = _decode_text(data)
    raw = re.sub(r"\\'([0-9a-fA-F]{2})", " ", raw)
    raw = re.sub(r"\\par[d]?", "\n", raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", raw)
    raw = raw.replace("{", " ").replace("}", " ")
    raw = re.sub(r"\s+", " ", raw).strip()
    return raw


def extract_unique_parsing_keys_via_openai(*, positions: List[str], api_key: str) -> List[str]:
    items = [str(x).strip() for x in (positions or []) if str(x).strip()]
    if not items:
        return []
    key = (api_key or "").strip()
    if not key:
        return extract_parsing_keys_from_positions(items)

    system = (
        "Ты извлекаешь уникальные ключи (короткие поисковые фразы) из заявки. "
        "На входе список строк позиций/номенклатуры. "
        "На выходе верни ТОЛЬКО JSON-массив строк. "
        "Каждая строка: 2-6 слов, без количеств, цен, единиц измерения, размеров, артикулов, номеров, телефонов, email, URL. "
        "Дубликаты и почти-дубликаты объедини. "
        "Ключи должны подходить для запуска веб-парсинга поставщиков по номенклатуре."
    )
    user = "\n".join(items[:200])

    payload = {
        "model": os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
        "temperature": 0,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "response_format": {"type": "json_object"},
    }

    url = (os.getenv("OPENAI_BASE_URL") or "https://api.openai.com").rstrip("/") + "/v1/chat/completions"

    try:
        with httpx.Client(timeout=30.0) as client:
            r = client.post(url, headers={"Authorization": f"Bearer {key}"}, json=payload)
            r.raise_for_status()
            data = r.json() or {}
            content = (((data.get("choices") or [{}])[0].get("message") or {}).get("content") or "").strip()
    except Exception:
        return extract_parsing_keys_from_positions(items)

    if not content:
        return extract_parsing_keys_from_positions(items)

    arr = None
    try:
        obj = json.loads(content)
        if isinstance(obj, list):
            arr = obj
        elif isinstance(obj, dict):
            for k in ("keys", "keywords", "items"):
                v = obj.get(k)
                if isinstance(v, list):
                    arr = v
                    break
    except Exception:
        arr = None

    if not isinstance(arr, list):
        return extract_parsing_keys_from_positions(items)

    out: List[str] = []
    seen: set[str] = set()
    for it in arr:
        s = str(it or "").strip()
        if not s:
            continue
        s = " ".join(s.split())
        low = s.lower()
        if low in seen:
            continue
        seen.add(low)
        out.append(s)

    return out if out else extract_parsing_keys_from_positions(items)


def _extract_pdf_text(data: bytes) -> str:
    from io import BytesIO

    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    parts: List[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text:
            parts.append(page_text)
    return "\n".join(parts)


def _extract_pdf_tables_text(data: bytes) -> str:
    try:
        from io import BytesIO

        import pdfplumber

        parts: List[str] = []
        with pdfplumber.open(BytesIO(data)) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables() or []
                for table in tables:
                    for row in table or []:
                        cells = [str(c).strip() for c in (row or []) if c is not None and str(c).strip()]
                        if cells:
                            parts.append(" ".join(cells))
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_docling_text(filename: str, data: bytes) -> str:
    if not DOCLING_AVAILABLE:
        raise RecognitionDependencyError("Docling is not available")

    import os
    import tempfile

    suffix = os.path.splitext(filename or "document")[1] or ".pdf"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(data or b"")
        tmp_path = tmp.name
    try:
        converter = DocumentConverter()
        res = converter.convert(tmp_path)
        doc = res.document
        try:
            return doc.export_to_markdown() or ""
        except Exception:
            pass
        try:
            return doc.export_to_text() or ""
        except Exception:
            pass

        try:
            return doc.export_to_json() or ""
        except Exception:
            return ""
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


def _extract_docx_text(data: bytes) -> str:
    from io import BytesIO

    from docx import Document

    doc = Document(BytesIO(data))
    parts: List[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if (c.text or "").strip()]
            if cells:
                parts.append(" ".join(cells))
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _extract_xlsx_text(data: bytes) -> str:
    from io import BytesIO

    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(data), data_only=True)
    parts: List[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            row_parts: List[str] = []
            for cell in row:
                if cell is None:
                    continue
                s = str(cell).strip()
                if s:
                    row_parts.append(s)
            if row_parts:
                parts.append(" ".join(row_parts))
    return "\n".join(parts)


def _ocr_image_bytes(data: bytes) -> str:
    from io import BytesIO

    try:
        from PIL import Image
    except Exception as e:
        raise RecognitionDependencyError(f"Pillow is required for OCR: {e}")

    try:
        import pytesseract
    except Exception as e:
        raise RecognitionDependencyError(f"pytesseract is required for OCR: {e}")

    image = Image.open(BytesIO(data))
    return pytesseract.image_to_string(image, lang="rus+eng") or ""


def _ocr_via_parser_service(filename: str, content: bytes) -> str:
    try:
        with httpx.Client(timeout=30.0) as client:
            resp = client.post(
                "http://127.0.0.1:9004/ocr/extract-text",
                files={"file": (filename or "file", content)},
            )
        if resp.status_code != 200:
            return ""
        data = resp.json() or {}
        return (data.get("text") or "").strip()
    except Exception:
        return ""


def _ocr_pdf_bytes(data: bytes) -> str:
    try:
        import pypdfium2 as pdfium
    except Exception as e:
        raise RecognitionDependencyError(f"pypdfium2 is required for OCR PDF fallback: {e}")

    try:
        import pytesseract
    except Exception as e:
        raise RecognitionDependencyError(f"pytesseract is required for OCR PDF fallback: {e}")

    pdf = pdfium.PdfDocument(data)
    parts: List[str] = []
    for i in range(len(pdf)):
        page = pdf[i]
        pil_image = page.render(scale=2).to_pil()
        parts.append(pytesseract.image_to_string(pil_image, lang="rus+eng") or "")
    return "\n".join(parts)


def smart_recognize_positions_from_file(filename: str, content: bytes) -> List[str]:
    """
    Smart recognition that automatically selects the best engine.
    
    Priority:
    1. Digital PDF -> direct text extraction
    2. PaddleOCR for PDFs with tables
    3. EasyOCR for images and scanned PDFs
    4. Fallback to existing methods
    """
    name = (filename or "").lower()
    ext = name.split(".")[-1] if "." in name else ""

    # Images: OCR first (fast path for photos/scans)
    if ext in {"png", "jpg", "jpeg"}:
        try:
            text_content = _ocr_via_parser_service(filename, content)
            if text_content and text_content.strip():
                items = normalize_positions(parse_positions_from_text(text_content))
                needs_enrich = False
                if not items or len(items) < 12:
                    needs_enrich = True
                if not needs_enrich:
                    low = text_content.lower()
                    if ("труба" not in low) and ("гофр" not in low) and ("раструб" not in low):
                        needs_enrich = True
                if needs_enrich and EASYOCR_AVAILABLE:
                    try:
                        extra = smart_extract_text(filename, content)
                    except Exception:
                        extra = ""
                    if extra and extra.strip():
                        merged = (text_content.rstrip() + "\n" + extra.lstrip()).strip()
                        merged_items = normalize_positions(parse_positions_from_text(merged))
                        if merged_items and len(merged_items) > len(items or []):
                            return merged_items
                if items:
                    return items
        except Exception:
            pass

        if EASYOCR_AVAILABLE:
            try:
                text_content = smart_extract_text(filename, content)
                if text_content and text_content.strip():
                    return normalize_positions(parse_positions_from_text(text_content))
            except Exception:
                pass

        try:
            text_content = _ocr_image_bytes(content)
            if text_content and text_content.strip():
                return normalize_positions(parse_positions_from_text(text_content))
        except Exception:
            pass

    # 1) Prefer digital PDF text extraction when possible (fastest & best)
    if ext == "pdf":
        # 1a) Docling: best structured conversion (layout + tables + OCR when needed)
        if DOCLING_AVAILABLE:
            try:
                text = _extract_docling_text(filename, content)
                items = parse_positions_from_text(text)
                return normalize_positions(items)
            except Exception:
                pass

        # Use existing PDF text (no OCR) if it looks like a digital PDF
        try:
            pdf_text = _extract_pdf_text(content)
            if pdf_text and len(pdf_text.strip()) > 50:
                items = parse_positions_from_text(pdf_text)
                items = normalize_positions(items)
                if items:
                    return items
        except Exception:
            pass

    # Try PaddleOCR for PDFs first
    if ext == "pdf" and PADDLEOCR_AVAILABLE:
        try:
            # Check if digital PDF
            if paddle_is_digital_pdf(content):
                # Use existing digital extraction
                text_content = _extract_pdf_tables_text(content)
                if not text_content:
                    text_content = _extract_pdf_text(content)
            else:
                # Use PaddleOCR for scanned PDFs
                tables = extract_tables_from_pdf(content)
                text_content = tables_to_text(tables)
            
            if text_content and text_content.strip():
                return normalize_positions(parse_positions_from_text(text_content))
        except Exception as e:
            import logging
            logging.getLogger(__name__).warning(f"PaddleOCR failed: {e}, falling back")
    
    # Try EasyOCR for PDFs (and any remaining cases)
    if EASYOCR_AVAILABLE and name.endswith(".pdf"):
        try:
            text_content = smart_extract_text(filename, content)
            if text_content and text_content.strip():
                return normalize_positions(parse_positions_from_text(text_content))
        except Exception as e:
            import logging
            logging.getLogger(__name__).warning(f"EasyOCR failed: {e}, falling back")
    
    # Fallback to existing methods
    return recognize_positions_from_file(filename, content, RecognitionEngine.auto)


def recognize_positions_from_file(filename: str, content: bytes, engine: RecognitionEngine = RecognitionEngine.auto) -> List[str]:
    name = (filename or "").lower()

    text_content = ""

    if name.endswith(".pdf"):
        if engine in {RecognitionEngine.structured, RecognitionEngine.auto}:
            text_content = _extract_pdf_tables_text(content)
        if not text_content and engine == RecognitionEngine.structured:
            # Structured/table extraction can fail; fallback to raw text, then OCR.
            try:
                text_content = _extract_pdf_text(content)
            except Exception:
                text_content = ""
            if not (text_content or "").strip():
                text_content = _ocr_pdf_bytes(content)
        if not text_content and engine in {RecognitionEngine.auto, RecognitionEngine.ocr}:
            try:
                text_content = _extract_pdf_text(content)
            except Exception:
                text_content = ""
            if not (text_content or "").strip():
                text_content = _ocr_pdf_bytes(content)
        if engine == RecognitionEngine.docling:
            raise RecognitionDependencyError("Docling engine is not installed in this build")

    elif name.endswith(".docx"):
        text_content = _extract_docx_text(content)

    elif name.endswith(".xlsx"):
        text_content = _extract_xlsx_text(content)

    elif name.endswith(".png") or name.endswith(".jpg") or name.endswith(".jpeg"):
        text_content = _ocr_image_bytes(content)

    elif name.endswith(".html") or name.endswith(".htm"):
        text_content = _extract_html_text(content)

    elif name.endswith(".eml"):
        try:
            text_content = _extract_eml_text(content)
        except Exception:
            text_content = _decode_text(content)

    elif name.endswith(".rtf"):
        text_content = _extract_rtf_text(content)

    else:
        text_content = _decode_text(content)

    return normalize_positions(parse_positions_from_text(text_content))
